from flask import Flask, jsonify, request
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from flask_socketio import SocketIO, emit
import os
from .ai_service import analyze_resume_with_ai
import json
import fitz  # PyMuPDF
import docx  # python-docx
import io
import hashlib
import time
from backend.tasks import process_job_resumes
from datetime import datetime

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///resumes.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

CORS(app)  # This will enable CORS for all routes
socketio = SocketIO(app, cors_allowed_origins="*")

# Global variables for tracking job completion
job_completion_trackers = {}  # job_id -> {total_resumes, completed_resumes}

# --- Database Configuration ---
basedir = os.path.abspath(os.path.dirname(__file__))
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(basedir, 'resumes.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)

# --- Database Models ---
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    jobs = db.relationship('Job', backref='user', lazy=True)

class Job(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    description = db.Column(db.Text, nullable=False)
    resumes = db.relationship('Resume', backref='job', lazy=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(100), nullable=False)
    candidate_name = db.Column(db.String(120), nullable=True)
    content = db.Column(db.Text, nullable=False)
    content_hash = db.Column(db.String(64), nullable=False, index=True)
    analysis = db.Column(db.Text, nullable=True)
    job_id = db.Column(db.Integer, db.ForeignKey('job.id'), nullable=False)

    __table_args__ = (db.UniqueConstraint('job_id', 'filename', name='_job_filename_uc'),
                      db.UniqueConstraint('job_id', 'content_hash', name='_job_hash_uc'))

class Feedback(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    resume_id = db.Column(db.Integer, db.ForeignKey('resume.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    original_bucket = db.Column(db.String(50), nullable=False)
    suggested_bucket = db.Column(db.String(50), nullable=True)
    feedback_type = db.Column(db.String(20), nullable=False)  # 'override', 'correction', 'improvement'
    feedback_text = db.Column(db.Text, nullable=True)
    confidence_score = db.Column(db.Float, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relationships
    resume = db.relationship('Resume', backref='feedbacks')
    user = db.relationship('User', backref='feedbacks')

class BucketOverride(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    resume_id = db.Column(db.Integer, db.ForeignKey('resume.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    original_bucket = db.Column(db.String(50), nullable=False)
    new_bucket = db.Column(db.String(50), nullable=False)
    reason = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relationships
    resume = db.relationship('Resume', backref='bucket_overrides')
    user = db.relationship('User', backref='bucket_overrides')

# --- WebSocket Events ---
@socketio.on('connect')
def handle_connect():
    print('Client connected')
    emit('status', {'message': 'Connected to server'})

@socketio.on('disconnect')
def handle_disconnect():
    print('Client disconnected')

def emit_progress_update(job_id, message, progress_type='info'):
    """Emit progress updates to connected clients"""
    socketio.emit('progress_update', {
        'job_id': job_id,
        'message': message,
        'type': progress_type,
        'timestamp': time.time()
    })

def check_job_completion(job_id):
    """Check if all resumes for a job are complete and emit completion event"""
    if job_id in job_completion_trackers:
        tracker = job_completion_trackers[job_id]
        tracker['completed_resumes'] += 1
        
        if tracker['completed_resumes'] >= tracker['total_resumes']:
            emit_progress_update(job_id, "All resumes processed successfully!", 'complete')
            # Clean up tracker
            del job_completion_trackers[job_id]

def process_resume_with_progress(job_id, resume_file, job_description):
    """Process a single resume with progress updates"""
    try:
        emit_progress_update(job_id, f"Processing {resume_file.filename}...", 'processing')
        
        content = ""
        filename = resume_file.filename
        file_stream = resume_file.read()

        if filename.endswith('.pdf'):
            emit_progress_update(job_id, f"Reading PDF: {filename}", 'info')
            pdf_doc = fitz.open(stream=file_stream, filetype='pdf')
            for page in pdf_doc:
                content += page.get_text()
            pdf_doc.close()
        elif filename.endswith('.docx'):
            emit_progress_update(job_id, f"Reading DOCX: {filename}", 'info')
            doc = docx.Document(io.BytesIO(file_stream))
            for para in doc.paragraphs:
                content += para.text + '\n'
        else:
            content = file_stream.decode('utf-8')

        if not content.strip():
            emit_progress_update(job_id, f"Skipped {filename}: Empty or unreadable", 'warning')
            return None, 'File is empty or unreadable'

        # Check for duplicates
        content_hash = hashlib.sha256(content.encode('utf-8')).hexdigest()
        existing_by_hash = Resume.query.filter_by(job_id=job_id, content_hash=content_hash).first()
        if existing_by_hash:
            emit_progress_update(job_id, f"Skipped {filename}: Duplicate content", 'warning')
            return None, 'Duplicate content'

        emit_progress_update(job_id, f"Analyzing {filename} with AI...", 'processing')
        analysis_text = analyze_resume_with_ai(job_description, content)
        analysis_json = json.loads(analysis_text)
        
        # Handle case where AI returns "Name Not Found"
        candidate_name = analysis_json.get('candidate_name', 'Not Provided')
        if candidate_name.strip().lower() == 'name not found':
            candidate_name = 'Not Provided'

        emit_progress_update(job_id, f"Completed analysis for {filename}", 'success')
        
        # Check if all resumes for this job are complete
        check_job_completion(job_id)
        
        return {
            'filename': filename,
            'candidate_name': candidate_name,
            'content': content,
            'content_hash': content_hash,
            'analysis': analysis_text
        }, None
        
    except Exception as e:
        emit_progress_update(job_id, f"Error processing {resume_file.filename}: {str(e)}", 'error')
        
        # Still check completion even on error
        check_job_completion(job_id)
        return None, str(e)

@app.route('/api/analyze', methods=['POST'])
def analyze_resumes():
    # --- Temp: Get or create a default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        default_user = User(username='default_user')
        db.session.add(default_user)
        db.session.commit()
    # --- End Temp ---

    if 'jobDescription' not in request.form:
        return jsonify({'error': 'No job description provided'}), 400

    job_description = request.form['jobDescription']
    resumes = request.files.getlist('resumes')

    # Check if a job with this description already exists FOR THIS USER.
    job = Job.query.filter_by(description=job_description, user_id=default_user.id).first()

    # If it doesn't exist, create a new one for this user.
    if not job:
        job = Job(description=job_description, user_id=default_user.id)
        db.session.add(job)
        db.session.flush()  # Use flush to get the job.id before committing.
        db.session.commit()  # Commit the job to the database

    emit_progress_update(job.id, f"Preparing {len(resumes)} resumes for background processing...", 'start')

    # Prepare resume data for background processing
    resumes_data = []
    for resume_file in resumes:
        if resume_file:
            try:
                content = ""
                filename = resume_file.filename
                file_stream = resume_file.read()

                if filename.endswith('.pdf'):
                    pdf_doc = fitz.open(stream=file_stream, filetype='pdf')
                    for page in pdf_doc:
                        content += page.get_text()
                    pdf_doc.close()
                elif filename.endswith('.docx'):
                    doc = docx.Document(io.BytesIO(file_stream))
                    for para in doc.paragraphs:
                        content += para.text + '\n'
                else:
                    content = file_stream.decode('utf-8')

                if content.strip():
                    resumes_data.append({
                        'filename': filename,
                        'content': content
                    })
                else:
                    emit_progress_update(job.id, f"Skipped {filename}: Empty or unreadable", 'warning')
                    
            except Exception as e:
                emit_progress_update(job.id, f"Error reading {resume_file.filename}: {str(e)}", 'error')

    if not resumes_data:
        emit_progress_update(job.id, "No valid resumes to process", 'warning')
        return jsonify({
            'message': 'No valid resumes to process',
            'job_id': job.id,
            'processed_files': [],
            'skipped_files': [{'filename': f.filename, 'reason': 'Invalid file'} for f in resumes if f]
        })

    # Initialize completion tracker for this job
    job_completion_trackers[job.id] = {
        'total_resumes': len(resumes_data),
        'completed_resumes': 0
    }

    # Queue background job using Celery
    process_job_resumes.delay(job.id, resumes_data, job_description)

    return jsonify({
        'message': f'Queued {len(resumes_data)} resumes for background processing',
        'job_id': job.id,
        'status': 'queued',
        'total_resumes': len(resumes_data)
    })

@app.route('/api/jobs', methods=['GET'])
def get_jobs():
    """Returns a list of all jobs for the current user."""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify([]) # No user, no jobs
    # --- End Temp ---

    jobs = Job.query.filter_by(user_id=default_user.id).order_by(Job.id.desc()).all()
    return jsonify([
        {
            'id': job.id, 
            'description': job.description,
            'resume_count': len(job.resumes)
        } for job in jobs
    ])

@app.route('/api/jobs/<int:job_id>', methods=['GET'])
def get_job_details(job_id):
    """Returns details and resumes for a specific job, checking user ownership."""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    job = Job.query.filter_by(id=job_id, user_id=default_user.id).first_or_404()
    
    resumes_data = [
        {
            'id': resume.id,
            'filename': resume.filename,
            'candidate_name': resume.candidate_name,
            'analysis': json.loads(resume.analysis) if resume.analysis else None
        } 
        for resume in job.resumes
    ]
    return jsonify({
        'id': job.id,
        'description': job.description,
        'resumes': resumes_data
    })

@app.route('/api/jobs/<int:job_id>', methods=['DELETE'])
def delete_job(job_id):
    """Delete a job and all its associated resumes"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    job = Job.query.filter_by(id=job_id, user_id=default_user.id).first()
    if not job:
        return jsonify({'error': 'Job not found'}), 404

    try:
        # Get resume count for confirmation
        resume_count = len(job.resumes)
        
        # Delete all associated resumes first (cascade)
        for resume in job.resumes:
            db.session.delete(resume)
        
        # Delete the job
        db.session.delete(job)
        db.session.commit()
        
        return jsonify({
            'message': f'Job "{job.description[:50]}..." deleted successfully',
            'deleted_resumes': resume_count,
            'job_id': job_id
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to delete job: {str(e)}'}), 500

@app.route('/api/data')
def get_data():
    return jsonify({'message': 'Hello from the Flask backend!'})

# --- Feedback Loop API Endpoints ---

@app.route('/api/feedback', methods=['POST'])
def submit_feedback():
    """Submit feedback for a resume analysis"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    data = request.get_json()
    
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    required_fields = ['resume_id', 'original_bucket', 'feedback_type']
    for field in required_fields:
        if field not in data:
            return jsonify({'error': f'Missing required field: {field}'}), 400
    
    # Verify resume exists and belongs to user
    resume = Resume.query.filter_by(id=data['resume_id']).first()
    if not resume:
        return jsonify({'error': 'Resume not found'}), 404
    
    # Check if resume belongs to user's job
    if resume.job.user_id != default_user.id:
        return jsonify({'error': 'Unauthorized access to resume'}), 403
    
    try:
        feedback = Feedback(
            resume_id=data['resume_id'],
            user_id=default_user.id,
            original_bucket=data['original_bucket'],
            suggested_bucket=data.get('suggested_bucket'),
            feedback_type=data['feedback_type'],
            feedback_text=data.get('feedback_text'),
            confidence_score=data.get('confidence_score')
        )
        
        db.session.add(feedback)
        db.session.commit()
        
        return jsonify({
            'message': 'Feedback submitted successfully',
            'feedback_id': feedback.id
        }), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to submit feedback: {str(e)}'}), 500

@app.route('/api/feedback/<int:resume_id>', methods=['GET'])
def get_resume_feedback(resume_id):
    """Get all feedback for a specific resume"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    # Verify resume exists and belongs to user
    resume = Resume.query.filter_by(id=resume_id).first()
    if not resume:
        return jsonify({'error': 'Resume not found'}), 404
    
    if resume.job.user_id != default_user.id:
        return jsonify({'error': 'Unauthorized access to resume'}), 403
    
    feedbacks = Feedback.query.filter_by(resume_id=resume_id).order_by(Feedback.created_at.desc()).all()
    
    return jsonify([{
        'id': f.id,
        'original_bucket': f.original_bucket,
        'suggested_bucket': f.suggested_bucket,
        'feedback_type': f.feedback_type,
        'feedback_text': f.feedback_text,
        'confidence_score': f.confidence_score,
        'created_at': f.created_at.isoformat()
    } for f in feedbacks])

@app.route('/api/override', methods=['POST'])
def submit_override():
    """Submit a bucket override for a resume"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    data = request.get_json()
    
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    required_fields = ['resume_id', 'original_bucket', 'new_bucket', 'reason']
    for field in required_fields:
        if field not in data:
            return jsonify({'error': f'Missing required field: {field}'}), 400
    
    # Verify resume exists
    resume = Resume.query.get(data['resume_id'])
    if not resume:
        return jsonify({'error': 'Resume not found'}), 404
    
    try:
        # Create override record
        override = BucketOverride(
            resume_id=data['resume_id'],
            user_id=default_user.id,  # --- Temp: Use default user ---
            original_bucket=data['original_bucket'],
            new_bucket=data['new_bucket'],
            reason=data['reason'],
            timestamp=datetime.utcnow()
        )
        
        db.session.add(override)
        db.session.commit()
        
        return jsonify({'message': 'Override submitted successfully', 'override_id': override.id}), 201
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to save override: {str(e)}'}), 500

@app.route('/api/override/<int:resume_id>', methods=['GET'])
def get_resume_overrides(resume_id):
    """Get all bucket overrides for a specific resume"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    # Verify resume exists and belongs to user
    resume = Resume.query.filter_by(id=resume_id).first()
    if not resume:
        return jsonify({'error': 'Resume not found'}), 404
    
    if resume.job.user_id != default_user.id:
        return jsonify({'error': 'Unauthorized access to resume'}), 403
    
    overrides = BucketOverride.query.filter_by(resume_id=resume_id).order_by(BucketOverride.created_at.desc()).all()
    
    return jsonify([{
        'id': o.id,
        'original_bucket': o.original_bucket,
        'new_bucket': o.new_bucket,
        'reason': o.reason,
        'created_at': o.created_at.isoformat()
    } for o in overrides])

@app.route('/api/feedback/stats', methods=['GET'])
def get_feedback_stats():
    """Get feedback statistics for the current user"""
    # --- Temp: Use default user ---
    default_user = User.query.filter_by(username='default_user').first()
    if not default_user:
        return jsonify({'error': 'User not found'}), 404
    # --- End Temp ---

    try:
        # Get total feedback count
        total_feedback = Feedback.query.filter_by(user_id=default_user.id).count()
        
        # Get feedback by type
        feedback_by_type = db.session.query(
            Feedback.feedback_type,
            db.func.count(Feedback.id)
        ).filter_by(user_id=default_user.id).group_by(Feedback.feedback_type).all()
        
        # Get bucket override count
        total_overrides = BucketOverride.query.filter_by(user_id=default_user.id).count()
        
        # Get most common original buckets that get overridden
        common_overrides = db.session.query(
            BucketOverride.original_bucket,
            db.func.count(BucketOverride.id)
        ).filter_by(user_id=default_user.id).group_by(BucketOverride.original_bucket).order_by(
            db.func.count(BucketOverride.id).desc()
        ).limit(5).all()
        
        return jsonify({
            'total_feedback': total_feedback,
            'feedback_by_type': dict(feedback_by_type),
            'total_overrides': total_overrides,
            'common_overrides': dict(common_overrides)
        })
        
    except Exception as e:
        return jsonify({'error': f'Failed to get feedback stats: {str(e)}'}), 500

@app.route('/')
def home():
    return "Hello from the Backend!" 